import io
from datetime import datetime, date
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.data.db import get_db
from app.data.orm import (
    Pedido, PedidoExterno, Autoparte, Inventario,
    Usuario, Cliente, EstadoPedido, Categoria
)
from app.security.auth import verify_api_key

router = APIRouter(prefix="/v1/reportes", tags=["Reportes"])

FORMATOS = {"pdf", "xlsx", "docx"}


# ── Helpers generadores ────────────────────────────────────────────────

def _pdf(title: str, headers: list, rows: list) -> io.BytesIO:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(letter), leftMargin=0.5 * inch, rightMargin=0.5 * inch)
    styles = getSampleStyleSheet()
    elems = [
        Paragraph(title, styles["Title"]),
        Paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]),
        Spacer(1, 20),
    ]
    data = [headers] + rows
    col_count = len(headers)
    col_width = (10 * inch) / col_count
    tbl = Table(data, colWidths=[col_width] * col_count, repeatRows=1)
    tbl.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2FF")]),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("PADDING", (0, 0), (-1, -1), 4),
        ])
    )
    elems.append(tbl)
    doc.build(elems)
    buf.seek(0)
    return buf


def _xlsx(title: str, headers: list, rows: list) -> io.BytesIO:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]

    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    alt_fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin

    for i, row in enumerate(rows):
        ws.append([str(v) if v is not None else "" for v in row])
        fill = alt_fill if i % 2 == 1 else None
        for cell in ws[i + 2]:
            cell.border = thin
            cell.alignment = Alignment(horizontal="center")
            if fill:
                cell.fill = fill

    for col in ws.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _docx(title: str, headers: list, rows: list) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1E3A5F")
        shading.set(qn("w:val"), "clear")
        hdr_cells[i].paragraphs[0]._p.get_or_add_pPr().append(shading)

    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val) if val is not None else ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _respond(buf: io.BytesIO, fmt: str, filename: str) -> StreamingResponse:
    media = {
        "pdf": "application/pdf",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    return StreamingResponse(
        buf,
        media_type=media[fmt],
        headers={"Content-Disposition": f'attachment; filename="{filename}.{fmt}"'},
    )


# ── Endpoints de reportes ─────────────────────────────────────────────

@router.get("/ventas")
async def reporte_ventas(
    formato: str = Query("pdf", description="pdf | xlsx | docx"),
    fecha_inicio: Optional[date] = Query(None),
    fecha_fin: Optional[date] = Query(None),
    db: Session = Depends(get_db),
    _: str = Depends(verify_api_key),
):
    """Reporte de ventas totales (pedidos internos + externos)"""
    if formato not in FORMATOS:
        raise HTTPException(status_code=400, detail="Formato inválido. Use: pdf, xlsx, docx")

    q_int = db.query(Pedido)
    q_ext = db.query(PedidoExterno)
    if fecha_inicio:
        q_int = q_int.filter(Pedido.fecha_pedido >= datetime.combine(fecha_inicio, datetime.min.time()))
        q_ext = q_ext.filter(PedidoExterno.fecha_pedido >= datetime.combine(fecha_inicio, datetime.min.time()))
    if fecha_fin:
        q_int = q_int.filter(Pedido.fecha_pedido <= datetime.combine(fecha_fin, datetime.max.time()))
        q_ext = q_ext.filter(PedidoExterno.fecha_pedido <= datetime.combine(fecha_fin, datetime.max.time()))

    pedidos_int = q_int.all()
    pedidos_ext = q_ext.all()

    headers = ["ID", "Tipo", "Cliente/Usuario", "Fecha", "Estado", "Total ($)"]
    rows = []
    for p in pedidos_int:
        rows.append([
            f"INT-{p.id}", "Interno",
            p.usuario.nombre if p.usuario else "N/A",
            p.fecha_pedido.strftime("%d/%m/%Y") if p.fecha_pedido else "N/A",
            p.estado.nombre if p.estado else "N/A",
            f"{float(p.total):.2f}",
        ])
    for p in pedidos_ext:
        rows.append([
            f"EXT-{p.id}", "Externo",
            p.cliente.nombre if p.cliente else "N/A",
            p.fecha_pedido.strftime("%d/%m/%Y") if p.fecha_pedido else "N/A",
            p.estado.nombre if p.estado else "N/A",
            f"{float(p.total):.2f}",
        ])

    total_ventas = sum(float(p.total) for p in pedidos_int) + sum(float(p.total) for p in pedidos_ext)
    rows.append(["", "", "", "", "TOTAL", f"{total_ventas:.2f}"])

    title = f"Reporte de Ventas - MACUIN"
    if fecha_inicio or fecha_fin:
        title += f" ({fecha_inicio or ''} al {fecha_fin or ''})"

    generators = {"pdf": _pdf, "xlsx": _xlsx, "docx": _docx}
    buf = generators[formato](title, headers, rows)
    return _respond(buf, formato, f"reporte_ventas_{datetime.now().strftime('%Y%m%d')}")


@router.get("/inventario")
async def reporte_inventario(
    formato: str = Query("pdf"),
    db: Session = Depends(get_db),
    _: str = Depends(verify_api_key),
):
    """Reporte de inventario actual"""
    if formato not in FORMATOS:
        raise HTTPException(status_code=400, detail="Formato inválido")

    items = db.query(Inventario).join(Autoparte).filter(Autoparte.activo == True).order_by(Autoparte.nombre).all()

    headers = ["ID", "Autoparte", "Categoría", "Marca", "Precio ($)", "Stock Actual", "Stock Mínimo", "Estado"]
    rows = []
    for inv in items:
        a = inv.autoparte
        estado = "BAJO STOCK" if inv.stock_actual <= inv.stock_minimo else "OK"
        rows.append([
            a.id, a.nombre,
            a.categoria.nombre if a.categoria else "N/A",
            a.marca or "N/A",
            f"{float(a.precio):.2f}",
            inv.stock_actual, inv.stock_minimo, estado,
        ])

    generators = {"pdf": _pdf, "xlsx": _xlsx, "docx": _docx}
    buf = generators[formato]("Reporte de Inventario - MACUIN", headers, rows)
    return _respond(buf, formato, f"reporte_inventario_{datetime.now().strftime('%Y%m%d')}")


@router.get("/pedidos")
async def reporte_pedidos(
    formato: str = Query("pdf"),
    estado_nombre: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    _: str = Depends(verify_api_key),
):
    """Reporte de pedidos por estado"""
    if formato not in FORMATOS:
        raise HTTPException(status_code=400, detail="Formato inválido")

    q_int = db.query(Pedido)
    q_ext = db.query(PedidoExterno)
    if estado_nombre:
        est = db.query(EstadoPedido).filter(EstadoPedido.nombre.ilike(f"%{estado_nombre}%")).first()
        if est:
            q_int = q_int.filter(Pedido.estado_id == est.id)
            q_ext = q_ext.filter(PedidoExterno.estado_id == est.id)

    pedidos_int = q_int.order_by(Pedido.fecha_pedido.desc()).all()
    pedidos_ext = q_ext.order_by(PedidoExterno.fecha_pedido.desc()).all()

    headers = ["ID", "Tipo", "Cliente/Usuario", "Fecha", "Estado", "Artículos", "Total ($)"]
    rows = []
    for p in pedidos_int:
        rows.append([
            f"INT-{p.id}", "Interno",
            p.usuario.nombre if p.usuario else "N/A",
            p.fecha_pedido.strftime("%d/%m/%Y") if p.fecha_pedido else "N/A",
            p.estado.nombre if p.estado else "N/A",
            len(list(p.detalles)),
            f"{float(p.total):.2f}",
        ])
    for p in pedidos_ext:
        rows.append([
            f"EXT-{p.id}", "Externo",
            p.cliente.nombre if p.cliente else "N/A",
            p.fecha_pedido.strftime("%d/%m/%Y") if p.fecha_pedido else "N/A",
            p.estado.nombre if p.estado else "N/A",
            len(list(p.detalles)),
            f"{float(p.total):.2f}",
        ])

    generators = {"pdf": _pdf, "xlsx": _xlsx, "docx": _docx}
    buf = generators[formato]("Reporte de Pedidos - MACUIN", headers, rows)
    return _respond(buf, formato, f"reporte_pedidos_{datetime.now().strftime('%Y%m%d')}")


@router.get("/usuarios")
async def reporte_usuarios(
    formato: str = Query("pdf"),
    db: Session = Depends(get_db),
    _: str = Depends(verify_api_key),
):
    """Reporte de usuarios internos y clientes externos"""
    if formato not in FORMATOS:
        raise HTTPException(status_code=400, detail="Formato inválido")

    usuarios = db.query(Usuario).order_by(Usuario.nombre).all()
    clientes = db.query(Cliente).order_by(Cliente.nombre).all()

    headers = ["ID", "Tipo", "Nombre", "Email", "Rol/Info", "Registro", "Estado"]
    rows = []
    for u in usuarios:
        pedidos_count = db.query(Pedido).filter(Pedido.usuario_id == u.id).count()
        rows.append([
            f"U-{u.id}", "Interno",
            u.nombre, u.email,
            f"{u.rol.nombre} | {pedidos_count} pedidos",
            u.fecha_registro.strftime("%d/%m/%Y") if u.fecha_registro else "N/A",
            "Activo" if u.activo else "Inactivo",
        ])
    for c in clientes:
        pedidos_count = db.query(PedidoExterno).filter(PedidoExterno.cliente_id == c.id).count()
        rows.append([
            f"C-{c.id}", "Externo (Cliente)",
            c.nombre, c.email,
            f"Tel: {c.telefono or 'N/A'} | {pedidos_count} pedidos",
            c.fecha_registro.strftime("%d/%m/%Y") if c.fecha_registro else "N/A",
            "Activo" if c.activo else "Inactivo",
        ])

    generators = {"pdf": _pdf, "xlsx": _xlsx, "docx": _docx}
    buf = generators[formato]("Reporte de Usuarios - MACUIN", headers, rows)
    return _respond(buf, formato, f"reporte_usuarios_{datetime.now().strftime('%Y%m%d')}")


@router.get("/autopartes-mas-vendidas")
async def reporte_autopartes_mas_vendidas(
    formato: str = Query("pdf"),
    db: Session = Depends(get_db),
    _: str = Depends(verify_api_key),
):
    """Reporte de autopartes más vendidas"""
    if formato not in FORMATOS:
        raise HTTPException(status_code=400, detail="Formato inválido")

    from app.data.orm import DetallePedido, DetallePedidoExterno

    # Aggregate from internal + external orders
    int_q = (
        db.query(
            Autoparte.id.label("id"),
            Autoparte.nombre.label("nombre"),
            Autoparte.marca.label("marca"),
            func.sum(DetallePedido.cantidad).label("cantidad_vendida"),
            func.sum(DetallePedido.cantidad * DetallePedido.precio_unitario).label("ingresos"),
        )
        .join(DetallePedido, DetallePedido.autoparte_id == Autoparte.id)
        .group_by(Autoparte.id, Autoparte.nombre, Autoparte.marca)
        .all()
    )

    ext_q = (
        db.query(
            Autoparte.id.label("id"),
            func.sum(DetallePedidoExterno.cantidad).label("cantidad_vendida"),
            func.sum(DetallePedidoExterno.cantidad * DetallePedidoExterno.precio_unitario).label("ingresos"),
        )
        .join(DetallePedidoExterno, DetallePedidoExterno.autoparte_id == Autoparte.id)
        .group_by(Autoparte.id)
        .all()
    )

    # Combine
    ext_map = {r.id: (r.cantidad_vendida or 0, float(r.ingresos or 0)) for r in ext_q}
    combined = {}
    for r in int_q:
        ec, ei = ext_map.get(r.id, (0, 0.0))
        combined[r.id] = {
            "nombre": r.nombre, "marca": r.marca,
            "cantidad": (r.cantidad_vendida or 0) + ec,
            "ingresos": float(r.ingresos or 0) + ei,
        }
    for id_, (ec, ei) in ext_map.items():
        if id_ not in combined:
            a = db.query(Autoparte).filter(Autoparte.id == id_).first()
            combined[id_] = {"nombre": a.nombre if a else "?", "marca": a.marca if a else "?",
                             "cantidad": ec, "ingresos": ei}

    sorted_items = sorted(combined.values(), key=lambda x: x["cantidad"], reverse=True)

    headers = ["Autoparte", "Marca", "Cantidad Vendida", "Ingresos Totales ($)"]
    rows = [[v["nombre"], v["marca"] or "N/A", v["cantidad"], f"{v['ingresos']:.2f}"] for v in sorted_items]

    generators = {"pdf": _pdf, "xlsx": _xlsx, "docx": _docx}
    buf = generators[formato]("Autopartes Más Vendidas - MACUIN", headers, rows)
    return _respond(buf, formato, f"reporte_mas_vendidas_{datetime.now().strftime('%Y%m%d')}")
